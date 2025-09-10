import os
import mimetypes
from flask import Flask, render_template, request, send_file
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
import google.generativeai as genai
from datetime import datetime
from docx import Document
from reportlab.pdfgen import canvas
from dotenv import load_dotenv

# ------------------------
# Load environment variables
# ------------------------
load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# ------------------------
# Flask Configuration
# ------------------------
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = "uploads"
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# ------------------------
# Database Model
# ------------------------
class Meeting(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(200))
    transcript = db.Column(db.Text)
    notes = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# create tables safely inside app context
with app.app_context():
    db.create_all()

# ------------------------
# Gemini Processing
# ------------------------
model = genai.GenerativeModel("gemini-2.0-flash")

def process_audio_with_gemini(filepath):
    # Detect mime type (mp3, wav, ogg, opus etc.)
    mime_type, _ = mimetypes.guess_type(filepath)
    if mime_type is None:
        mime_type = "audio/wav"  # fallback

    # Read file as bytes
    with open(filepath, "rb") as f:
        audio_bytes = f.read()

    # Send audio + instructions to Gemini
    response = model.generate_content(
        [
            "Transcribe this meeting audio. If not in English, translate to English. "
            "Then generate structured meeting notes including: Summary, Key Points, "
            "Decisions, Action Items, and Sentiment. Return JSON format.",
            {
                "mime_type": mime_type,
                "data": audio_bytes
            }
        ]
    )
    return response.text

# ------------------------
# Routes
# ------------------------
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        if "file" not in request.files:
            return "No file uploaded!"
        file = request.files["file"]
        if file.filename == "":
            return "No selected file!"
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # Gemini processing
        notes = process_audio_with_gemini(filepath)

        # Save to DB
        meeting = Meeting(filename=filename, transcript="(Transcript handled by Gemini)", notes=notes)
        db.session.add(meeting)
        db.session.commit()

        return render_template("notes.html", notes=notes, meeting=meeting)
    return render_template("index.html")

@app.route("/history")
def history():
    meetings = Meeting.query.order_by(Meeting.created_at.desc()).all()
    return render_template("history.html", meetings=meetings)

@app.route("/download/<int:id>/<string:format>")
def download(id, format):
    meeting = Meeting.query.get_or_404(id)

    if format == "word":
        filepath = f"outputs/notes_{id}.docx"
        doc = Document()
        doc.add_heading("Meeting Notes", 0)
        doc.add_paragraph(meeting.notes)
        doc.save(filepath)
        return send_file(filepath, as_attachment=True)

    elif format == "pdf":
        filepath = f"outputs/notes_{id}.pdf"
        c = canvas.Canvas(filepath)
        c.drawString(100, 800, "Meeting Notes")
        text = c.beginText(100, 780)
        for line in meeting.notes.splitlines():
            text.textLine(line)
        c.drawText(text)
        c.save()
        return send_file(filepath, as_attachment=True)

    return "Invalid format!"

# ------------------------
# Run the App
# ------------------------
if __name__ == "__main__":
    if not os.path.exists("uploads"):
        os.makedirs("uploads")
    if not os.path.exists("outputs"):
        os.makedirs("outputs")
    app.run(debug=True)
